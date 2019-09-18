# -*- coding: utf-8 -*-
"""
AE Detaction Prediction

Created on 2019

@Author: Deng, Wentao <wdeng@cn.imshealth.com>  
@Supervisor； Zhang, Yao <YZhang03@cn.imshealth.com>           
@Cleaner: Zhu, Lexi <yingyan.zhu@cn.imshealth.com> and
          Deng, Wentao <wdeng@cn.imshealth.com> 
"""
import sys, os
sys.path.append(os.path.abspath('.'))

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re
import numpy as np
import logging
import pandas as pd

from fuzzywuzzy import fuzz, process
from keras.models import model_from_json, load_model

from ensemble.bert_start import bc
import tensorflow as tf
glob = tf.get_default_graph()

class CNN_LSTM_Predictor(object):
    '''
    CNN_LSTM_Predictor relying on pretrained CNN_BiLSTM model can help to detect
    potential adverse events in input interview records. 
    '''
    logger = logging.getLogger(__name__)

    logger.info('default reference folder directory')
    main_dir = '.'
    # initial
    def __init__(self, m_json = None, m_weight=None, regex=None, **kwargs):
        self.model = load_model(self.main_dir +'/BERT_CNN_LSTM/model_0905.h5')
        if 'domain' in kwargs:
            if m_json == None:
                m_json = self.main_dir + '/BERT_CNN_LSTM/'+ kwargs['domain'] + '/model_leukemia_0428.json'
            if m_weight == None:
                m_weight = self.main_dir + '/BERT_CNN_LSTM/'+ kwargs['domain'] + '/model_leukemia_0428.h5'
            
            self.m_json = m_json
            self.m_weight = m_weight

            json_file = open(self.m_json, 'r')
            load_model_json = json_file.read()
            json_file.close()
            model = model_from_json(load_model_json)
            model.load_weights(self.m_weight)
            model.compile(optimizer = 'adam', 
                        loss = 'categorical_crossentropy', metrics = ['accuracy'])
            self.model = model
        
        if regex == None:
            regex = self.main_dir + '/AE_oncology_CLEAN.txt'
        self.regex = regex

    # format dialogues
    def dialogues(self, filepath):
        '''
        Format sentences into completed Q&A dialogues or pairwise dialogues
        '''
        contents = Document(filepath)
        con = {}
        para_idx = {}
        self.logger.info('Starting format sentences...')
        for i, para in enumerate(contents.paragraphs):
            text = re.sub(' ', '', para.text)
            if text != "":
                if text[0] not in 'QAMI123456789' or (i == 0 and text[0] in 'A1'):
                    k = 0
                else:
                    if text[0] in 'QM' or (k == 0 and text[0] in'A1'): 
                        k = i 
                        key = i
                        para_idx[key] = str(i)
                        con[key] = text
                    else:
                        para_idx[key] += ' ' + str(i)
                        con[key] += text
        return para_idx, con

    # highlight and color the predicted positives
    def forecast(self, contents):
        '''
        Detecting AEs with in docx file and return docx instances.
        '''
        para_idx, con = self.dialogues(contents)

        with open(self.regex, 'r', encoding = 'utf-8') as reg:
            lines = [line.strip() for line in reg.readlines() if line.strip() !=""]
            reg.close()
        # highlight
        for key in con:
            vector = [c for c in bc([con[key]])]
            vec_m = np.reshape(vector, (1, 24, 64)) 
            model = self.model
            with glob.as_default():
                proba_pred = model.predict(vec_m)
            y_test_pred = np.argmax(proba_pred, axis = 1)

            # catch positive words
            if y_test_pred == 0:
                res = process.extractOne(con[key], lines, scorer = fuzz.partial_ratio)
                y_test_pred = 1 if res[1] >= 85 else 0

            # catch negative words
            if len(con[key]) < 50 and ('不担心' in con[key] or '不明显' in con[key]): 
                y_test_pred = 0
            
            # save docx 
            if y_test_pred:
                idx = [int(num) for num in para_idx[key].split(' ')]
                for i in idx:
                    para = contents.paragraphs[i]
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        return contents
        
    def output_df(self, contents):
        '''
        output dataframe containing labels(1/0), contents, pargraph indices.
        '''
        para_idx, con = self.dialogues(contents)
        with open(self.regex, 'r', encoding = 'utf-8') as reg:
            lines = [line.strip() for line in reg.readlines() if line.strip() !=""]
            reg.close()
        
        texts = [con[key] for key in con]
        vecs = [c for c in bc(texts)]
        vecs = np.reshape(vecs, (-1,24,64))

        model = self.model
 
        with glob.as_default():
            proba_pred = model.predict(vecs)
    
        # store labels
        labels = np.argmax(proba_pred, axis=1)
        for i, text in enumerate(texts):
            if labels[i] == 0:
                res = process.extractOne(text, lines, scorer = fuzz.partial_ratio)
                labels[i] = 1 if res[1] >= 85 else 0
            if len(text)<50 and ('不担心' in text or '不明显' in text):
                labels[i] = 0
            

        para_num = [para_idx[k] for k in para_idx]
        cont = [con[k] for k in con]
        return pd.DataFrame({ 'contents': cont, \
                              'paragraph_idx': para_num,\
                              'labels': labels})
    
    def tableMap2Docx(self, contents, tab_output):
        '''
        After manually checking tables online, this function can map target labels 
        into paragraphs accordingly.
        contents: docx instances
        tab_output： dataframe output after submitting
        '''
        cons = Document(contents)
        self.logger.debug(f'tab_output feature name {tab_output.columns}')
        for lab, par in tab_output[['labels','paragraph_idx']].values:
            if lab:
                idxs = [int(n) for n in par.split(' ')]
                for i in idxs:
                        para = cons.paragraphs[i]
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
          
        return cons
        

if __name__ == '__main__':
    logging.basicConfig(level=logging.DEBUG, format='%(asctime)s %(name)-12s %(levelname)-8s %(funcName)s() %(message)s')
    ### provide filepath
    basicDir = '.'
    filepath = basicDir + '/A.docx'
    ## default run oncology, if leukemia, run CNN_LSTM_Predictor(domain = 'leukemia')
    pred = CNN_LSTM_Predictor()
    # doc = pred.forecast(filepath)
    # ## output directory
    # doc.save('./pred_' + os.path.basename(filepath))
    import time
    tic = time.time()
    df = pred.output_df(filepath)
    toc = time.time()
    print(f'Time elapsed: {toc-tic}')
    print(df)
## 6 cpus/64 GB memory:44s
## 4 cpus/64 GB memory: 4.5mins
## shutdown BertService use the code as follow in your terminal:
## bert-serving-terminate -port 5555