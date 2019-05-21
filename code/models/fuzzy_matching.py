import docx 
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

import re
import glob

import docx 
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

import re
import glob

# segmentation
import jieba_fast as jieba




def dialogues(filepath):
    contents = Document(filepath)
    con = {}
    para_idx ={}
    for i, para in enumerate(contents.paragraphs):
        if para.text!="":
            if para.text[0] not in 'QAMI1234567890':
                k = 0
            else:
                if para.text[0] in 'QM' or (k==0 and para.text[0] =='A'):
                    k = i
                    key = 'QM'+str(i)
                    para_idx[key] = str(i)
                    con[key] = para.text
                else:
                    para_idx[key] += ' '+str(i)
                    con[key] +=para.text
    return para_idx, con
  

def match(names, filepath):
    contents = Document(filepath)
    para_idx, con = dialogues(filepath)
    with open(names,'r', encoding='utf-8') as p:
        names = p.readlines()
        for key in con:
            for name in names:
                name = name.strip()
                reg ='(.*?)'+ name +'(.*?).*'
                match = re.search(reg, con[key], re.I)
                if match:
                    idx = [int(n) for n in para_idx[key].split(' ')]
                    for i in idx:
                        for run in contents.paragraphs[i].runs:
                            run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        p.close()
    return contents

path = 'model/data/Dictionary/AE_oncology_CLEAN.txt'
filepaths = glob.glob('model/data/label_unlabel_data/reviewed/reviewed_no_selected/*.docx')

doc = match(path, filepaths[5])
doc.save('model/a.docx')

#with open(path, 'r', encoding='utf-8') as p:
#  names =p.readlines()
#  for name in names:
#    print(name.strip())