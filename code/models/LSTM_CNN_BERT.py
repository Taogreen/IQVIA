import keras
from keras.preprocessing.sequence import pad_sequences
from keras.models import Sequential,Model, model_from_json
from keras.layers import *
from keras import metrics
from keras.optimizers import *

from keras import backend as K
import tensorflow as tf

import numpy as np
np.random.seed(100) # for replication

# load positive and negative sample
x_pos = np.load('model/data/BERT_output/bert_pos_reviewed_0415.npy')
y_pos = np.ones(len(x_pos), dtype = int)

x_neg = np.load('model/data/BERT_output/bert_neg_reviewed_0415.npy')
y_neg = np.zeros(len(x_neg), dtype = int)

x = np.concatenate([x_pos, x_neg])
y = np.concatenate([y_pos, y_neg])
# shuffle data
index = np.arange(len(y))
np.random.shuffle(index)
x,y = x[index], y[index]


from sklearn.model_selection import train_test_split
x_train, x_test, y_train, y_test = train_test_split(x, y, test_size = 0.1, random_state = 100, stratify=y)

from imblearn.over_sampling import SMOTE
sm = SMOTE(random_state=2)
x_train, y_train  = sm.fit_sample(x_train, y_train)
X_train = np.reshape(x_train,(len(x_train), 24, 64))
X_test = np.reshape(x_test, (len(x_test),24, 64))
from keras.utils import to_categorical
Y_train = to_categorical(y_train)
Y_test = to_categorical(y_test)

###################
NO TRAIN/TEST SPLIT
###################

from imblearn.over_sampling import SMOTE
sm = SMOTE(random_state=2)
x_train, y_train  = sm.fit_sample(x, y)
X_train = np.reshape(x_train,(len(x_train), 24, 64))
#X_test = np.reshape(x_test, (len(x_test),24, 32))
from keras.utils import to_categorical
Y_train = to_categorical(y_train)
#Y_test = to_categorical(y_test)

###################################################################

inputs = Input(shape = (24,64), name = 'input')
conv1_1 = Conv1D(128, kernel_size=2, padding = 'valid', strides = 1)(inputs)
#conv1_2 = Conv1D(128, kernel_size=2, padding = 'valid', strides = 1)(conv1_1)
activation1 = Activation('relu')(conv1_1)
maxpool1 = MaxPool1D(3)(activation1)

conv2_1 = Conv1D(128, kernel_size=3, padding = 'valid', strides = 1)(inputs)
#conv2_2 = Conv1D(128, kernel_size=3, padding = 'valid', strides = 1)(conv2_1)
activation2 = Activation('relu')(conv2_1)
maxpool2 = MaxPool1D(3)(activation2)

conv3_1 = Conv1D(128, kernel_size=4, padding = 'valid', strides = 1)(inputs)
#conv3_2 = Conv1D(128, kernel_size=4, padding = 'valid', strides = 1)(conv3_1)
activation3 = Activation('relu')(conv3_1)
maxpool3 = MaxPool1D(3)(activation3)

concat = Concatenate(axis = 1)([maxpool1, maxpool2, maxpool3])
dropout = Dropout(0.6)(concat)
#flatten = Flatten()(dropout)
lstm_1 = Bidirectional(LSTM(100), name = 'bilstm')(dropout)
fc1 = Dense(200, activation='softmax', name = 'fc1')(lstm_1)
attention = multiply([lstm_1,fc1], name='attention_mul')
#dropout = Dropout(0.5)(attention)
output = Dense(2, activation = 'softmax', name='output')(attention)
m = Model(input=inputs, output = output)
m.summary()
m.compile(optimizer='adam', 
              loss='categorical_crossentropy', metrics=['accuracy'])
m.fit(X_train, Y_train,  batch_size = 100, epochs = 20)

proba_test = m.predict(X_test)
y_test_pred = np.argmax(proba_test, axis=1)

from sklearn.metrics import confusion_matrix, classification_report
confusion_matrix(y_test,y_test_pred)
print(classification_report(y_test,y_test_pred))

x_new=np.reshape(x,[len(x),24,64])
proba = m.predict(x_new)
y_pred = np.argmax(proba, axis=1)

from sklearn.metrics import confusion_matrix, classification_report
confusion_matrix(y,y_pred)
print(classification_report(y,y_pred))

##########################################
########## check false negative ##########
##########################################
ind = [i for i in range(len(x)) if y[i]==1 and y_pred[i]==0]
sorted(index[ind])

####################### training set probability ##########################
#proba_train = m.predict(X_train)
#threshold = np.linspace(0,1, num=100)
#precision = []
#recall = []
#for val in threshold:
#  y_train_pred=[1 if p1>val else 0 for _,p1 in proba_train]
#  TP = sum([1 if v1==1 and v2==1 else 0 for v1, v2 in zip(y_train, y_train_pred)])
#  pos_pred = sum(y_train_pred)
#  pos_train = sum(y_train)
#  if pos_pred>0:
#    precision.append(TP/pos_pred)
#    recall.append(TP/pos_train)
#
#np.save('model/data/recall_precision_proba/proba', proba_train)
#np.save('model/data/recall_precision_proba/recall', recall)
#np.save('model/data/recall_precision_proba/precision',precision)


# save model
model_json = m.to_json()
with open('model/data/BERT_CNN_LSTM/model_selected_data_0415.json', 'w') as json:
  json.write(model_json)
m.save_weights('model/data/BERT_CNN_LSTM/model_selected_data_0415.h5')

#json_file = open('model/data/BERT_CNN_LSTM/model_selected_data_0409.json','r')
#load_model_json = json_file.read()
#json_file.close()
#model = model_from_json(load_model_json)
#model.load_weights('model/data/BERT_CNN_LSTM/model.h5')
#
#model.compile(optimizer='adam', 
#              loss='categorical_crossentropy', metrics=['accuracy'])
#m.fit(X_train, Y_train, validation_data=(X_test, Y_test), batch_size = 100, epochs = 10)

proba_pred = model.predict(X_test)
y_test_pred = np.argmax(proba_pred, axis = 1)
from sklearn.metrics import confusion_matrix, classification_report
confusion_matrix(y_test, y_test_pred)
print(classification_report(y_test,y_test_pred))

# LSTM alone : 0.57    0.58    0.57 
# LSTM+attention: 0.43      0.66      0.52      
# CNN: 0.68 0.61 0.64
#      0.67 0.56 0.61
#      0.57 0.69 0.62
#      0.69 0.61 0.65
#      0.58 0.63 0.61
#      0.65 0.62 0.63
#      0.58 0.63 0.61
#      0.68 0.58 0.63
#      0.38 0.69 0.49
#      0.76 0.59 0.67

# LSTM： p = [.41,.26,.43,.40,.50,.42,.32,.35,.36,.31]
#        r = [.66,.72,.75,.65,.55,.62,.66,.70,.62,.69]
#        f = [.51,.38,.54,.50,.52,.50,.43,.47,.46,.43]

# C_BiALSTM： p = [.61,.71,.44,.42,.42,.66,.56,.48,.51,.67]
#             r = [.59,.56,.65,.68,.62,.58,.61,.68,.63,.61]
#             f = [.60,.63,.52,.52,.50,.62,.58,.56,.56,.64]
np.save('model/data/BiALSTM-predict', proba_pred)
np.save('model/data/BiALSTM-Original', y_test)

#################################################
#################### 白血病 ######################
#################################################
np.random.seed(100) # for replication

# load positive and negative sample
x_pos = np.load('model/data/BERT_output/leukemia/bert_pos_0422.npy')
y_pos = np.ones(len(x_pos), dtype = int)

x_neg = np.load('model/data/BERT_output/leukemia/bert_neg_0422.npy')
y_neg = np.zeros(len(x_neg), dtype = int)

x = np.concatenate([x_pos, x_neg])
y = np.concatenate([y_pos, y_neg])
# shuffle data
index = np.arange(len(y))
np.random.shuffle(index)
x,y = x[index], y[index]


from sklearn.model_selection import train_test_split
x_train, x_test, y_train, y_test = train_test_split(x, y, test_size = 0.1, random_state = 100, stratify=y)

from imblearn.over_sampling import SMOTE
sm = SMOTE(random_state=2)
x_train, y_train  = sm.fit_sample(x_train, y_train)
X_train = np.reshape(x_train,(len(x_train), 24, 64))
X_test = np.reshape(x_test, (len(x_test),24, 64))
from keras.utils import to_categorical
Y_train = to_categorical(y_train)
Y_test = to_categorical(y_test)

############################################
########## NO TRAIN/TEST SPLIT #############
############################################
from imblearn.over_sampling import SMOTE
sm = SMOTE(random_state=2)
x_train, y_train  = sm.fit_sample(x, y)
X_train = np.reshape(x_train,(len(x_train), 24, 64))
#X_test = np.reshape(x_test, (len(x_test),24, 32))
from keras.utils import to_categorical
Y_train = to_categorical(y_train)
#Y_test = to_categorical(y_test)

inputs = Input(shape = (24,64), name = 'input')
conv1_1 = Conv1D(128, kernel_size=2, padding = 'valid', strides = 1)(inputs)
#conv1_2 = Conv1D(128, kernel_size=2, padding = 'valid', strides = 1)(conv1_1)
activation1 = Activation('relu')(conv1_1)
maxpool1 = MaxPool1D(3)(activation1)

conv2_1 = Conv1D(128, kernel_size=3, padding = 'valid', strides = 1)(inputs)
#conv2_2 = Conv1D(128, kernel_size=3, padding = 'valid', strides = 1)(conv2_1)
activation2 = Activation('relu')(conv2_1)
maxpool2 = MaxPool1D(3)(activation2)

conv3_1 = Conv1D(128, kernel_size=4, padding = 'valid', strides = 1)(inputs)
#conv3_2 = Conv1D(128, kernel_size=4, padding = 'valid', strides = 1)(conv3_1)
activation3 = Activation('relu')(conv3_1)
maxpool3 = MaxPool1D(3)(activation3)

concat = Concatenate(axis = 1)([maxpool1, maxpool2, maxpool3])
dropout = Dropout(0.6)(concat)
#flatten = Flatten()(dropout)
lstm_1 = Bidirectional(LSTM(100), name = 'bilstm')(dropout)
fc1 = Dense(200, activation='softmax', name = 'fc1')(lstm_1)
attention = multiply([lstm_1,fc1], name='attention_mul')
#dropout = Dropout(0.5)(attention)
output = Dense(2, activation = 'softmax', name='output')(attention)
m = Model(input=inputs, output = output)
m.summary()
m.compile(optimizer='adam', 
              loss='categorical_crossentropy', metrics=['accuracy'])
m.fit(X_train, Y_train,  batch_size = 100, epochs = 20)

proba_test = m.predict(X_test)
y_test_pred = np.argmax(proba_test, axis=1)

from sklearn.metrics import confusion_matrix, classification_report
confusion_matrix(y_test,y_test_pred)
print(classification_report(y_test,y_test_pred))

x_new=np.reshape(x,[len(x),24,64])
proba = m.predict(x_new)
y_pred = np.argmax(proba, axis=1)

from sklearn.metrics import confusion_matrix, classification_report
confusion_matrix(y,y_pred)
print(classification_report(y,y_pred))

##########################################
########## check false negative ##########
##########################################
ind = [i for i in range(len(x)) if y[i]==1 and y_pred[i]==0]
sorted(index[ind])

model_json = m.to_json()
with open('model/data/BERT_CNN_LSTM/model_leukemia_0428.json', 'w') as json:
  json.write(model_json)
m.save_weights('model/data/BERT_CNN_LSTM/model_leukemia_0428.h5')


#################################################
#################### Prediction #################
#################################################
from keras.models import model_from_json

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import glob, os, re
import numpy as np

#from bert_serving.server import get_args_parser
#from bert_serving.server import BertServer
#import sys
#args = get_args_parser().parse_args(['-model_dir','model/data/chinese_L-12_H-768_A-12',
#                                     '-num_worker','2',
#                                     '-pooling_strategy','REDUCE_MEAN_MAX',
#                                     '-max_seq_len','350'
#                                    ])
#server = BertServer(args)
#server.start()
from bert_serving.client import BertClient
bc = BertClient()


class prediction:
  def __init__(self, filepath, model, regex = 'model/data/fuzzy_match/leukemia.txt'):
    self.filepath = filepath
    self.model = model
    self.regex = regex
  
  # format dialogues
  def dialogues(self):
    contents = Document(self.filepath)
    con = {}
    para_idx ={}
    for i, para in enumerate(contents.paragraphs):
      text = re.sub(' ','',para.text)
      if text!="":
        if text[0] not in 'QAMI123456789' or (i==0 and text[0] in 'A1'):
          k = 0
        else:
          if text[0] in 'QM' or (k==0 and text[0] in'A1'):
            k = i
            key = 'QM'+str(i)
            para_idx[key] = str(i)
            con[key] = text
          else:
            para_idx[key] += ' '+str(i)
            con[key] +=text
    return para_idx, con
  # highlight and color the predicted positives
  def highlight(self):
    contents = Document(self.filepath)
    para_idx, con = self.dialogues()
    with open(self.regex,'r') as reg:
      lines = reg.readlines()
      reg.close()
      
    # highlight
    for key in con:
      vector = bc.encode([con[key]])
      vec_m = np.reshape(vector, (1, 24, 64))
      proba_pred = model.predict(vec_m)
      y_test_pred = np.argmax(proba_pred, axis=1)
#      if y_test_pred == 0:
#        for line in lines:
#          line = line.strip()
#          match = re.search(line, con[key])
#          if match:
#            y_test_pred =1
#            break   
      if len(con[key])<50 and ('不担心' in con[key] or '不明显' in con[key]):
        y_test_pred = 0
      if y_test_pred:
        idx = [int(num) for num in para_idx[key].split(' ')]
        for i in idx:
          para = contents.paragraphs[i]
          for run in para.runs:
            run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    return contents
  
model_json = 'model/data/BERT_CNN_LSTM/model_leukemia_0428.json'
model_weights = 'model/data/BERT_CNN_LSTM//model_leukemia_0428.h5'
    
# load model and its weights
json_file = open(model_json,'r')
load_model_json = json_file.read()
json_file.close()

model = model_from_json(load_model_json)
model.load_weights(model_weights)
model.compile(optimizer='adam', 
              loss='categorical_crossentropy', metrics=['accuracy'])
path = glob.glob('model/data/label_unlabel_data/predict/pred_0409_test/*.docx')
for p in path:
  c = prediction(p, model)
  doc = c.highlight()
  doc.save('model/data/label_unlabel_data/predict/pred_0409_test_result/'+os.path.basename(p))

## print probability
#p1 = 'model/data/label_unlabel_data/predict/pred_0409_test/18CR412-CML 市场研究 – 访问笔录-河南省人们医院陈玉清主任医师.docx'
#c = prediction(p1, model)
#para_idx, con = c.dialogues()
#for key in con:
#  vector = bc.encode([con[key]])
#  vec_m = np.reshape(vector, (1, 24, 64))
#  proba_pred = model.predict(vec_m)
#  y_test_pred = np.argmax(proba_pred, axis = 1)
#  print (proba_pred, ' ', con[key])
#doc = c.highlight()
#doc.save('model/data/label_unlabel_data/predict/'+os.path.basename(p1))

#contents = Document(p1)
#con = {}
#para_idx ={}
#for i, para in enumerate(contents.paragraphs):
#   if para.text!="":
#      if para.text[0] not in 'QAMI123456789' or (i==0 and para.text[0] in 'A1'):
#          k = 0
#      else:
#          if para.text[0] in 'QM' or (k==0 and para.text[0] in'A1'):
#            k = i
#            key = 'QM'+str(i)
#            para_idx[key] = str(i)
#            con[key] = para.text
#          else:
#            para_idx[key] += ' '+str(i)
#            con[key] +=para.text
#import re
#regex = open('model/data/fuzzy_match/global.txt','r')
#a = 'Q：总体，那就是说一些抗感染的。A：都算上。'
#
#for line in regex.readlines():
#  line = line.strip()
#  match = re.search(line, a)
#  if match:
#    print(line)
#    break   
    
